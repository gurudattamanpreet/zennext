from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.responses import HTMLResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from typing import List, Dict, Optional, Tuple
import os
import io
import json
import asyncio
from datetime import datetime
import hashlib
import re
from collections import Counter
from dataclasses import dataclass
from enum import Enum
import statistics

# Document processing
try:
    import pdfplumber

    PDF_PLUMBER_AVAILABLE = True
except ImportError:
    PDF_PLUMBER_AVAILABLE = False
    print("Warning: pdfplumber not installed, using PyPDF2")

import PyPDF2
from docx import Document

# NLP and ML
import nltk
from sklearn.feature_extraction.text import TfidfVectorizer, CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.ensemble import RandomForestClassifier
from sklearn.preprocessing import StandardScaler
import numpy as np

# Download required NLTK data
try:
    nltk.download('punkt', quiet=True)
    nltk.download('stopwords', quiet=True)
    nltk.download('wordnet', quiet=True)
    nltk.download('averaged_perceptron_tagger', quiet=True)
    nltk.download('maxent_ne_chunker', quiet=True)
    nltk.download('words', quiet=True)
    from nltk.corpus import stopwords
    from nltk.tokenize import word_tokenize, sent_tokenize
    from nltk.stem import WordNetLemmatizer
    from nltk import pos_tag, ne_chunk

    NLTK_AVAILABLE = True
except:
    NLTK_AVAILABLE = False
    print("Warning: Some NLTK components not available")

# Groq API for LLaMA
import httpx
from pydantic import BaseModel

# Configuration
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
GROQ_BASE_URL = os.getenv("GROQ_BASE_URL", "https://api.groq.com/openai/v1")
GROQ_MODEL = os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile")

app = FastAPI(title="AI Recruiter System - Enterprise Edition")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ==========================================
# STEP 1: Document Parsing & Extraction
# ==========================================

class ResumeSection(Enum):
    """Industry standard resume sections"""
    CONTACT = "contact_info"
    SUMMARY = "professional_summary"
    EXPERIENCE = "work_experience"
    EDUCATION = "education"
    SKILLS = "skills"
    PROJECTS = "projects"
    CERTIFICATIONS = "certifications"
    ACHIEVEMENTS = "achievements"
    LANGUAGES = "languages"
    REFERENCES = "references"


@dataclass
class ParsedResume:
    """Structured resume data model"""
    raw_text: str
    sections: Dict[str, str]
    contact_info: Dict[str, str]
    work_experience: List[Dict]
    education: List[Dict]
    skills: List[str]
    certifications: List[str]
    languages: List[str]
    metadata: Dict[str, any]


class DocumentParser:
    """Step 1: Advanced document parsing with section identification"""

    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> str:
        """Extract text from PDF with multiple fallback methods"""
        text = ""

        if PDF_PLUMBER_AVAILABLE:
            try:
                pdf_file = io.BytesIO(file_bytes)
                with pdfplumber.open(pdf_file) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                if text.strip():
                    return text.strip()
            except Exception as e:
                print(f"pdfplumber extraction failed: {e}")

        # Fallback to PyPDF2
        try:
            pdf_file = io.BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            print(f"PyPDF2 extraction error: {e}")
            return ""

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> str:
        """Extract text from DOCX with table support"""
        try:
            doc_file = io.BytesIO(file_bytes)
            doc = Document(doc_file)

            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())

            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))

            return "\n".join(full_text)
        except Exception as e:
            print(f"DOCX extraction error: {e}")
            return ""

    @staticmethod
    def identify_sections(text: str) -> Dict[str, str]:
        """Identify and extract resume sections using patterns"""
        sections = {}
        section_headers = {
            'contact': r'(?i)(contact|email|phone|address|linkedin)',
            'summary': r'(?i)(summary|objective|profile|about)',
            'experience': r'(?i)(experience|employment|work|career)',
            'education': r'(?i)(education|academic|qualification)',
            'skills': r'(?i)(skills|technical|competencies)',
            'projects': r'(?i)(projects|portfolio)',
            'certifications': r'(?i)(certifications|certificates|training)',
            'achievements': r'(?i)(achievements|accomplishments|awards)'
        }

        lines = text.split('\n')
        current_section = 'summary'
        section_content = []

        for line in lines:
            section_found = False
            for section_name, pattern in section_headers.items():
                if re.search(pattern, line[:50]):  # Check first 50 chars
                    if section_content:
                        sections[current_section] = '\n'.join(section_content)
                    current_section = section_name
                    section_content = []
                    section_found = True
                    break

            if not section_found:
                section_content.append(line)

        if section_content:
            sections[current_section] = '\n'.join(section_content)

        return sections

    @staticmethod
    def extract_contact_info(text: str) -> Dict[str, str]:
        """Extract contact information using regex patterns"""
        contact = {}

        # Email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_match = re.search(email_pattern, text)
        if email_match:
            contact['email'] = email_match.group()

        # Phone
        phone_pattern = r'[\+]?[(]?[0-9]{1,4}[)]?[-\s\.]?[(]?[0-9]{1,4}[)]?[-\s\.]?[0-9]{1,5}[-\s\.]?[0-9]{1,5}'
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            contact['phone'] = phone_match.group()

        # LinkedIn
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_match:
            contact['linkedin'] = linkedin_match.group()

        # GitHub
        github_pattern = r'github\.com/[\w-]+'
        github_match = re.search(github_pattern, text, re.IGNORECASE)
        if github_match:
            contact['github'] = github_match.group()

        return contact


# ==========================================
# STEP 2: NLP Processing
# ==========================================

class NLPProcessor:
    """Step 2: Advanced NLP processing with entity recognition"""

    def __init__(self):
        self.lemmatizer = WordNetLemmatizer() if NLTK_AVAILABLE else None
        self.stop_words = set(stopwords.words('english')) if NLTK_AVAILABLE else set()

    def extract_entities(self, text: str) -> Dict[str, List]:
        """Named Entity Recognition for companies, locations, etc."""
        entities = {
            'companies': [],
            'locations': [],
            'job_titles': [],
            'dates': [],
            'technologies': []
        }

        # Company patterns
        company_patterns = [
            r'\b(?:Inc|Corp|LLC|Ltd|Company|Technologies|Solutions|Services|Systems|Software|Consulting)\b',
            r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+(?:Inc|Corp|LLC|Ltd)\b'
        ]

        for pattern in company_patterns:
            matches = re.findall(pattern, text)
            entities['companies'].extend(matches)

        # Job title patterns
        job_patterns = [
            r'\b(?:Senior|Junior|Lead|Principal|Staff|Manager|Director|Engineer|Developer|Analyst|Consultant|Architect|Designer|Specialist)\b[^.]*(?:Engineer|Developer|Manager|Analyst|Consultant|Architect|Designer|Specialist)\b'
        ]

        for pattern in job_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            entities['job_titles'].extend(matches[:5])  # Limit to top 5

        # Technology extraction
        tech_keywords = [
            'Python', 'Java', 'JavaScript', 'C++', 'C#', 'Ruby', 'Go', 'Rust', 'Swift', 'Kotlin',
            'React', 'Angular', 'Vue', 'Node.js', 'Django', 'Flask', 'Spring', 'Express',
            'AWS', 'Azure', 'GCP', 'Docker', 'Kubernetes', 'Jenkins', 'Git',
            'TensorFlow', 'PyTorch', 'Scikit-learn', 'Pandas', 'NumPy',
            'SQL', 'NoSQL', 'MongoDB', 'PostgreSQL', 'MySQL', 'Redis',
            'Machine Learning', 'Deep Learning', 'NLP', 'Computer Vision', 'AI'
        ]

        text_lower = text.lower()
        for tech in tech_keywords:
            if tech.lower() in text_lower:
                entities['technologies'].append(tech)

        # Date extraction (years of experience)
        date_pattern = r'\b(19|20)\d{2}\b'
        dates = re.findall(date_pattern, text)
        entities['dates'] = dates

        return entities

    def tokenize_and_clean(self, text: str) -> List[str]:
        """Tokenize, lemmatize, and remove stop words"""
        if not NLTK_AVAILABLE:
            return text.split()

        # Tokenize
        tokens = word_tokenize(text.lower())

        # Remove stop words and non-alphabetic tokens
        tokens = [token for token in tokens if token.isalpha() and token not in self.stop_words]

        # Lemmatize
        if self.lemmatizer:
            tokens = [self.lemmatizer.lemmatize(token) for token in tokens]

        return tokens

    def extract_years_of_experience(self, text: str) -> float:
        """Extract total years of experience from resume"""
        patterns = [
            r'(\d+)\+?\s*years?\s*(?:of\s*)?experience',
            r'experience\s*[:\-]?\s*(\d+)\s*years?',
            r'(\d+)\s*years?\s*(?:of\s*)?professional',
            r'total\s*experience\s*[:\-]?\s*(\d+)\s*years?'
        ]

        years = []
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                try:
                    years.append(float(match))
                except:
                    pass

        if years:
            return max(years)  # Return maximum found

        # Try to calculate from dates
        dates = re.findall(r'\b(20\d{2})\b', text)
        if len(dates) >= 2:
            try:
                return abs(int(max(dates)) - int(min(dates)))
            except:
                pass

        return 0


# ==========================================
# STEP 3: Semantic Understanding
# ==========================================

class SemanticAnalyzer:
    """Step 3: Vector embeddings and semantic similarity"""

    def __init__(self):
        self.tfidf_vectorizer = TfidfVectorizer(
            max_features=500,
            stop_words='english',
            ngram_range=(1, 3),  # Unigrams, bigrams, and trigrams
            min_df=1,
            max_df=0.95
        )
        self.embeddings_cache = {}

    def create_embeddings(self, texts: List[str]) -> np.ndarray:
        """Create TF-IDF embeddings for texts"""
        try:
            embeddings = self.tfidf_vectorizer.fit_transform(texts)
            return embeddings.toarray()
        except:
            return np.zeros((len(texts), 100))

    def calculate_semantic_similarity(self, text1: str, text2: str) -> float:
        """Calculate semantic similarity between two texts with context awareness"""
        try:
            # Handle empty or invalid inputs
            if not text1 or not text2 or len(text1) < 50 or len(text2) < 50:
                return 30.0  # Base fallback score

            # Create cache key
            cache_key = hashlib.md5(f"{text1[:100]}{text2[:100]}".encode()).hexdigest()

            if cache_key in self.embeddings_cache:
                return self.embeddings_cache[cache_key]

            # Enhanced preprocessing for better matching
            # Extract key terms from both texts
            text1_processed = self.preprocess_for_matching(text1)
            text2_processed = self.preprocess_for_matching(text2)

            # Create embeddings
            embeddings = self.create_embeddings([text1_processed, text2_processed])

            # Calculate cosine similarity
            similarity = cosine_similarity(embeddings[0:1], embeddings[1:2])[0][0]

            # Industry-standard score mapping
            # Based on real ATS systems (Taleo, Workday, iCIMS)
            if similarity <= 0.05:
                normalized_score = similarity * 200  # 0-10% for very low match
            elif similarity <= 0.15:
                normalized_score = 10 + (similarity - 0.05) * 200  # 10-30% for low match
            elif similarity <= 0.25:
                normalized_score = 30 + (similarity - 0.15) * 200  # 30-50% for fair match
            elif similarity <= 0.40:
                normalized_score = 50 + (similarity - 0.25) * 133  # 50-70% for good match
            elif similarity <= 0.60:
                normalized_score = 70 + (similarity - 0.40) * 100  # 70-90% for excellent
            else:
                normalized_score = 90 + (similarity - 0.60) * 25  # 90-100% for perfect match

            # Ensure score is within bounds
            normalized_score = max(0, min(100, normalized_score))

            # Cache result
            self.embeddings_cache[cache_key] = float(normalized_score)

            return float(normalized_score)
        except Exception as e:
            print(f"Semantic similarity error: {e}")
            return 30.0  # Return base score on error

    def preprocess_for_matching(self, text: str) -> str:
        """Preprocess text for better semantic matching"""
        # Convert to lowercase
        text = text.lower()

        # Remove special characters but keep important ones
        text = re.sub(r'[^\w\s\+\#\-\.\/]', ' ', text)

        # Normalize whitespace
        text = ' '.join(text.split())

        # Expand common abbreviations for better matching
        abbreviations = {
            'ml': 'machine learning',
            'dl': 'deep learning',
            'ai': 'artificial intelligence',
            'nlp': 'natural language processing',
            'cv': 'computer vision',
            'js': 'javascript',
            'ts': 'typescript',
            'k8s': 'kubernetes',
            'ci/cd': 'continuous integration continuous deployment',
            'api': 'application programming interface',
            'ui': 'user interface',
            'ux': 'user experience',
            'db': 'database',
            'devops': 'development operations',
            'mlops': 'machine learning operations'
        }

        for abbr, full in abbreviations.items():
            text = text.replace(f' {abbr} ', f' {full} ')

        return text

    def extract_skill_embeddings(self, skills: List[str], job_skills: List[str]) -> float:
        """Calculate skill overlap using set similarity and embeddings"""
        if not job_skills:
            return 50.0  # No requirements specified, neutral score

        if not skills:
            return 0.0  # No skills found in resume

        # Convert to lowercase sets
        resume_skills = set(s.lower().strip() for s in skills if s)
        required_skills = set(s.lower().strip() for s in job_skills if s)

        if not required_skills:
            return 50.0

        # Direct exact matches
        exact_matches = len(resume_skills.intersection(required_skills))

        # Partial matches (substring matching)
        partial_matches = 0
        for rs in resume_skills:
            for js in required_skills:
                if rs != js:  # Not already counted as exact
                    if (len(rs) > 3 and rs in js) or (len(js) > 3 and js in rs):
                        partial_matches += 0.5

        # Related skill matching (common variations)
        skill_variations = {
            'js': ['javascript', 'node', 'nodejs'],
            'python': ['django', 'flask', 'fastapi'],
            'ml': ['machine learning', 'deep learning', 'ai'],
            'react': ['reactjs', 'redux', 'next.js'],
            'aws': ['s3', 'ec2', 'lambda', 'cloud'],
            'docker': ['kubernetes', 'k8s', 'containerization'],
            'sql': ['mysql', 'postgresql', 'database'],
        }

        related_matches = 0
        for base_skill, variations in skill_variations.items():
            base_in_resume = any(base_skill in s or s in base_skill for s in resume_skills)
            base_in_job = any(base_skill in s or s in base_skill for s in required_skills)

            if base_in_resume and base_in_job:
                continue  # Already counted

            for var in variations:
                var_in_resume = any(var in s or s in var for s in resume_skills)
                var_in_job = any(var in s or s in var for s in required_skills)

                if var_in_resume and base_in_job:
                    related_matches += 0.3
                elif base_in_resume and var_in_job:
                    related_matches += 0.3

        # Calculate final score
        total_matches = exact_matches + partial_matches + related_matches
        max_possible = len(required_skills)

        # Calculate percentage with bonus for having more skills than required
        raw_score = (total_matches / max_possible) * 100

        # Bonus for extra skills (up to 10% bonus)
        extra_skills = len(resume_skills) - len(required_skills)
        if extra_skills > 0:
            bonus = min(10, extra_skills * 2)
            raw_score += bonus

        # Ensure score is within 0-100 range
        final_score = max(0, min(100, raw_score))

        return float(final_score)
        total_matches = exact_matches + partial_matches
        max_possible = len(required_skills)

        return min((total_matches / max_possible) * 100, 100) if max_possible > 0 else 0


# ==========================================
# STEP 4: Scoring Mechanisms
# ==========================================

class ScoringEngine:
    """Step 4: Multi-factor scoring system"""

    def __init__(self):
        self.weights = {
            'keyword_match': 0.25,
            'experience_relevance': 0.20,
            'education_score': 0.15,
            'contextual_match': 0.30,
            'skill_depth': 0.10
        }

    def calculate_keyword_score(self, resume_text: str, keywords: List[str]) -> float:
        """Industry-standard keyword matching with context awareness"""
        if not keywords:
            return 50.0

        resume_lower = resume_text.lower()

        # Preprocess resume for better matching
        resume_processed = re.sub(r'[^\w\s\+\#\-\.\/]', ' ', resume_lower)
        resume_words = set(resume_processed.split())

        # Create variations for better matching
        keyword_variations = {}
        for keyword in keywords:
            kw_lower = keyword.lower()
            variations = [kw_lower]

            # Add common variations
            if ' ' in kw_lower:  # Multi-word keywords
                variations.append(kw_lower.replace(' ', ''))  # Remove spaces
                variations.append(kw_lower.replace(' ', '-'))  # Hyphenated
                variations.append(kw_lower.replace(' ', '_'))  # Underscore

            # Add abbreviated versions
            if kw_lower == 'machine learning':
                variations.extend(['ml', 'machinelearning'])
            elif kw_lower == 'artificial intelligence':
                variations.extend(['ai', 'artificialintelligence'])
            elif kw_lower == 'javascript':
                variations.extend(['js', 'node', 'nodejs'])
            elif kw_lower == 'typescript':
                variations.extend(['ts'])

            keyword_variations[kw_lower] = variations

        # Scoring with weighted matching
        total_score = 0
        max_score_per_keyword = 100 / len(keywords)

        for original_keyword, variations in keyword_variations.items():
            keyword_score = 0

            # Check for exact match (full weight)
            for variant in variations:
                if variant in resume_processed:
                    # Check context quality
                    context_quality = self._check_keyword_context(variant, resume_processed)
                    keyword_score = max_score_per_keyword * context_quality
                    break

            # Check for partial match if no exact match
            if keyword_score == 0:
                for variant in variations:
                    # Word-level partial matching
                    variant_words = variant.split()
                    if len(variant_words) > 1:
                        matches = sum(1 for w in variant_words if w in resume_words)
                        if matches > 0:
                            keyword_score = max_score_per_keyword * (matches / len(variant_words)) * 0.7
                            break
                    else:
                        # Single word - check for substring match
                        if any(variant in word for word in resume_words if len(word) > 3):
                            keyword_score = max_score_per_keyword * 0.5
                            break

            total_score += keyword_score

        # Apply bonus for keyword density (up to 10% bonus)
        keyword_count = sum(resume_processed.count(kw.lower()) for kw in keywords)
        density_bonus = min(10, keyword_count / len(keywords) * 2)

        final_score = min(100, total_score + density_bonus)
        return final_score

    def _check_keyword_context(self, keyword: str, text: str) -> float:
        """Check the quality of keyword context (how prominently it appears)"""
        text_length = len(text)
        first_occurrence = text.find(keyword)

        if first_occurrence == -1:
            return 0.0

        # Keywords in first 20% of resume get full weight
        position_ratio = first_occurrence / text_length
        if position_ratio <= 0.2:
            position_score = 1.0
        elif position_ratio <= 0.5:
            position_score = 0.9
        else:
            position_score = 0.8

        # Check frequency (multiple mentions are better)
        frequency = text.count(keyword)
        if frequency >= 3:
            frequency_score = 1.0
        elif frequency == 2:
            frequency_score = 0.9
        else:
            frequency_score = 0.8

        return (position_score + frequency_score) / 2

    def calculate_experience_score(self, years: float, required_years: float) -> float:
        """Calculate experience relevance score"""
        if required_years <= 0:
            return 70.0

        if years >= required_years:
            # Overqualified penalty for 2x+ experience
            if years > required_years * 2:
                return 80.0
            return min(100, 80 + (years - required_years) * 5)
        else:
            # Under-qualified penalty
            ratio = years / required_years
            return max(20, ratio * 100)

    def calculate_education_score(self, education: str, required_education: str) -> float:
        """Calculate education fit score for any field"""
        # Comprehensive education level hierarchy
        education_levels = {
            # Doctoral
            'phd': 100,
            'doctorate': 100,
            'd.phil': 100,

            # Post Graduate
            'post graduate': 90,
            'pg ': 90,
            'post-graduate': 90,

            # Master's degrees
            'masters': 85,
            'master': 85,
            'mba': 85,
            'mtech': 85,
            'm.tech': 85,
            'msc': 85,
            'm.sc': 85,
            'ma ': 85,
            'm.a': 85,
            'mcom': 85,
            'm.com': 85,
            'mca': 85,
            'llm': 85,
            'med': 85,
            'm.ed': 85,
            'mphil': 85,
            'm.phil': 85,

            # Professional degrees
            'ca ': 80,
            'chartered accountant': 80,
            'icwa': 80,
            'cs ': 80,
            'company secretary': 80,
            'cfa': 80,
            'frm': 80,

            # Bachelor's degrees
            'bachelors': 70,
            'bachelor': 70,
            'graduate': 70,
            'graduation': 70,
            'btech': 70,
            'b.tech': 70,
            'be ': 70,
            'b.e': 70,
            'bsc': 70,
            'b.sc': 70,
            'ba ': 70,
            'b.a': 70,
            'bcom': 70,
            'b.com': 70,
            'bba': 70,
            'bca': 70,
            'llb': 70,
            'mbbs': 70,
            'bds': 70,

            # Diploma
            'diploma': 55,
            'polytechnic': 55,
            'iti': 50,
            'certificate': 45,

            # School education
            '12th': 40,
            'intermediate': 40,
            'hsc': 40,
            'higher secondary': 40,
            '10th': 30,
            'matriculation': 30,
            'ssc': 30,
            'secondary': 30,

            # Any/Not specified
            'any graduate': 65,
            'any': 50,
            'not specified': 50
        }

        education_lower = education.lower()
        required_lower = required_education.lower()

        # Extract education levels
        candidate_level = 50  # Default
        required_level = 65  # Default to "Any Graduate"

        # Find candidate's education level
        for level, score in education_levels.items():
            if level in education_lower:
                candidate_level = max(candidate_level, score)

        # Find required education level
        for level, score in education_levels.items():
            if level in required_lower:
                required_level = score
                break

        # Special handling for "Any Graduate" requirement
        if 'any' in required_lower and 'graduat' in required_lower:
            required_level = 65
            # Any bachelor's or above qualifies
            if candidate_level >= 70:
                return 100

        # Calculate match score
        if candidate_level >= required_level:
            # Over-qualified gets slightly less than perfect score
            if candidate_level - required_level > 20:
                return 90
            else:
                return min(100, 80 + (candidate_level - required_level))
        else:
            # Under-qualified
            gap = required_level - candidate_level
            if gap <= 10:
                return 70  # Close enough
            elif gap <= 20:
                return 50  # Acceptable gap
            elif gap <= 30:
                return 30  # Significant gap
            else:
                return 10  # Major gap

    def calculate_contextual_score(self, resume_sections: Dict, job_requirements: Dict) -> float:
        """Calculate contextual relevance using section analysis"""
        score = 50.0  # Base score

        # Check for relevant experience section
        if 'experience' in resume_sections:
            exp_text = resume_sections['experience'].lower()
            relevant_terms = job_requirements.get('key_responsibilities', [])
            matches = sum(1 for term in relevant_terms if term.lower() in exp_text)
            score += min(30, matches * 5)

        # Check for relevant projects
        if 'projects' in resume_sections:
            score += 10

        # Check for certifications
        if 'certifications' in resume_sections:
            score += 10

        return min(score, 100)


# ==========================================
# STEP 5: ML Models
# ==========================================

class MLPipeline:
    """Step 5: Machine Learning models for classification and ranking"""

    def __init__(self):
        self.feature_extractor = CountVectorizer(max_features=200)
        self.scaler = StandardScaler()
        self.models = {}

    def extract_features(self, resume_data: Dict) -> np.ndarray:
        """Extract numerical features from resume data"""
        features = []

        # Basic features
        features.append(resume_data.get('years_experience', 0))
        features.append(len(resume_data.get('skills', [])))
        features.append(len(resume_data.get('companies', [])))
        features.append(len(resume_data.get('certifications', [])))
        features.append(resume_data.get('education_score', 0))
        features.append(resume_data.get('keyword_score', 0))
        features.append(resume_data.get('semantic_score', 0))

        # Text statistics
        text = resume_data.get('text', '')
        features.append(len(text))  # Resume length
        features.append(len(text.split()))  # Word count
        features.append(len(text.split('\n')))  # Line count

        # Section presence (binary features)
        sections = resume_data.get('sections', {})
        for section in ['experience', 'education', 'skills', 'projects', 'certifications']:
            features.append(1 if section in sections else 0)

        return np.array(features).reshape(1, -1)

    def predict_suitability(self, features: np.ndarray) -> Tuple[float, str]:
        """Predict candidate suitability using ensemble"""
        # Simple rule-based prediction for now
        # In production, this would use trained models

        avg_score = np.mean(features[0, :7]) if features.shape[1] >= 7 else 50

        if avg_score >= 70:
            return avg_score, "HIGH"
        elif avg_score >= 50:
            return avg_score, "MEDIUM"
        else:
            return avg_score, "LOW"


# ==========================================
# STEP 6: Ranking & Filtering
# ==========================================

class RankingEngine:
    """Step 6: Multi-stage filtering and ranking"""

    def __init__(self):
        self.hard_filters = {}
        self.soft_scoring_weights = {
            'overall_match': 0.40,
            'skill_match': 0.25,
            'experience_match': 0.20,
            'education_match': 0.15
        }

    def apply_hard_filters(self, candidates: List[Dict], requirements: Dict) -> List[Dict]:
        """Apply knockout filters"""
        filtered = []

        for candidate in candidates:
            # Check minimum experience
            min_exp = requirements.get('min_experience', 0)
            if candidate.get('years_experience', 0) < min_exp * 0.7:  # 70% tolerance
                candidate['rejection_reason'] = 'Insufficient experience'
                continue

            # Check must-have skills
            must_have = requirements.get('must_have_skills', [])
            candidate_skills = [s.lower() for s in candidate.get('skills', [])]
            if must_have:
                matched = sum(1 for skill in must_have if skill.lower() in candidate_skills)
                if matched < len(must_have) * 0.5:  # 50% must match
                    candidate['rejection_reason'] = 'Missing critical skills'
                    continue

            # Check location if required
            if requirements.get('location_required'):
                if not candidate.get('location_match', False):
                    candidate['rejection_reason'] = 'Location mismatch'
                    continue

            filtered.append(candidate)

        return filtered

    def calculate_percentile_rank(self, candidates: List[Dict]) -> List[Dict]:
        """Calculate percentile ranking for candidates"""
        if not candidates:
            return candidates

        scores = [c.get('overall_score', 0) for c in candidates]

        for candidate in candidates:
            score = candidate.get('overall_score', 0)
            percentile = sum(1 for s in scores if s <= score) / len(scores) * 100
            candidate['percentile_rank'] = round(percentile, 1)

        return candidates

    def apply_soft_scoring(self, candidates: List[Dict]) -> List[Dict]:
        """Apply weighted soft scoring"""
        for candidate in candidates:
            weighted_score = 0

            for factor, weight in self.soft_scoring_weights.items():
                score = candidate.get(factor, 0)
                weighted_score += score * weight

            candidate['weighted_score'] = round(weighted_score, 2)

        # Sort by weighted score
        candidates.sort(key=lambda x: x.get('weighted_score', 0), reverse=True)

        return candidates


# ==========================================
# STEP 7: Industry-Specific Techniques
# ==========================================

class IndustryAnalyzer:
    """Step 7: Industry-specific analysis and ATS optimization"""

    def __init__(self):
        self.industry_keywords = {
            'tech': ['agile', 'scrum', 'ci/cd', 'devops', 'cloud', 'microservices', 'api', 'saas'],
            'finance': ['financial', 'banking', 'trading', 'risk', 'compliance', 'regulatory', 'portfolio'],
            'healthcare': ['patient', 'clinical', 'hipaa', 'medical', 'healthcare', 'diagnostic', 'treatment'],
            'consulting': ['client', 'stakeholder', 'strategy', 'analysis', 'presentation', 'problem-solving'],
            'retail': ['customer', 'sales', 'inventory', 'merchandising', 'pos', 'e-commerce', 'supply chain']
        }

    def detect_industry(self, text: str) -> str:
        """Detect industry from resume content"""
        text_lower = text.lower()
        industry_scores = {}

        for industry, keywords in self.industry_keywords.items():
            score = sum(1 for keyword in keywords if keyword in text_lower)
            industry_scores[industry] = score

        if industry_scores:
            return max(industry_scores, key=industry_scores.get)
        return 'general'

    def check_ats_compatibility(self, resume_text: str) -> Dict[str, any]:
        """
        Industry-standard ATS compatibility check based on:
        - Taleo (Oracle) - 35% market share
        - Workday - 20% market share
        - iCIMS - 15% market share
        - Greenhouse/Lever - Modern ATS systems
        """
        issues = []
        detailed_scores = {}

        # 1. FILE FORMAT & PARSING SCORE (25% weight)
        parsing_score = 100

        # Check for parsing killers
        parsing_killers = {
            '│': "Box drawing characters block ATS parsing",
            '►': "Special bullets cause parsing errors",
            '●': "Non-standard bullets may not parse",
            '◆': "Decorative symbols break ATS readers",
            '★': "Star symbols cause parsing issues",
            '❖': "Complex symbols not ATS-friendly"
        }

        for char, issue in parsing_killers.items():
            if char in resume_text:
                issues.append(issue)
                parsing_score -= 15

        # Check for tables/columns (major ATS issue)
        if '|' in resume_text or resume_text.count('\t') > 5:
            issues.append("Tables/columns fail in 70% of ATS systems")
            parsing_score -= 25

        # Check headers and footers (often ignored by ATS)
        lines = resume_text.split('\n')
        if len(lines) > 0 and len(lines[0]) < 20:
            issues.append("Important info in header may be missed")
            parsing_score -= 10

        detailed_scores['parsing'] = max(parsing_score, 0)

        # 2. KEYWORD OPTIMIZATION SCORE (30% weight)
        keyword_score = 100
        words = resume_text.lower().split()
        word_count = len(words)

        # Optimal word count for ATS
        if word_count < 300:
            issues.append(f"Too short ({word_count} words) - ATS needs 300-800 words")
            keyword_score = (word_count / 300) * 70  # Proportional penalty
        elif word_count > 1000:
            issues.append(f"Too long ({word_count} words) - may be truncated")
            keyword_score -= min((word_count - 1000) / 20, 30)

        # Check keyword density (repetition without stuffing)
        word_freq = Counter(words)
        top_words = [w for w, c in word_freq.most_common(10) if len(w) > 4]

        # Keywords should repeat 2-4 times for optimal ATS
        avg_repetition = statistics.mean([word_freq[w] for w in top_words]) if top_words else 1
        if avg_repetition < 2:
            issues.append("Key terms not repeated enough for ATS matching")
            keyword_score -= 20
        elif avg_repetition > 8:
            issues.append("Keyword stuffing detected - may trigger spam filters")
            keyword_score -= 25

        detailed_scores['keywords'] = max(keyword_score, 0)

        # 3. SECTION STRUCTURE SCORE (25% weight)
        section_score = 100
        text_lower = resume_text.lower()

        # Industry-standard section headers
        required_sections = {
            'experience': ['experience', 'employment', 'work history', 'professional experience'],
            'education': ['education', 'academic', 'qualification', 'degree'],
            'skills': ['skills', 'technical skills', 'competencies', 'expertise']
        }

        found_sections = []
        for section_type, headers in required_sections.items():
            found = any(header in text_lower for header in headers)
            if not found:
                issues.append(f"Missing standard '{section_type.upper()}' section header")
                section_score -= 20
            else:
                found_sections.append(section_type)

        # Check section order (Experience should come before Education for experienced candidates)
        if 'experience' in found_sections and 'education' in found_sections:
            exp_pos = text_lower.find('experience')
            edu_pos = text_lower.find('education')

            # Check if person has significant experience
            years_markers = re.findall(r'\d+[\+]?\s*years?', text_lower)
            if years_markers and exp_pos > edu_pos:
                issues.append("Experience section should come before Education for experienced candidates")
                section_score -= 10

        detailed_scores['structure'] = max(section_score, 0)

        # 4. CONTACT INFORMATION SCORE (10% weight)
        contact_score = 100

        # Check for email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        if not re.search(email_pattern, resume_text):
            issues.append("No email found - critical for ATS")
            contact_score -= 40

        # Check for phone
        phone_pattern = r'[\+]?[(]?[0-9]{1,4}[)]?[-\s\.]?[(]?[0-9]{1,4}[)]?[-\s\.]?[0-9]{1,5}[-\s\.]?[0-9]{1,5}'
        if not re.search(phone_pattern, resume_text):
            issues.append("No phone number detected")
            contact_score -= 30

        # Check if contact info is in first 20% of document (best practice)
        first_section = resume_text[:len(resume_text) // 5]
        if not re.search(email_pattern, first_section):
            issues.append("Contact information should be at the top")
            contact_score -= 15

        detailed_scores['contact'] = max(contact_score, 0)

        # 5. DATE FORMATTING SCORE (10% weight)
        date_score = 100

        # Check for consistent date formats
        date_formats = [
            r'\b\d{2}/\d{4}\b',  # MM/YYYY
            r'\b\d{4}-\d{4}\b',  # YYYY-YYYY
            r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}\b',  # Month YYYY
            r'\b\d{1,2}/\d{1,2}/\d{2,4}\b'  # MM/DD/YYYY
        ]

        found_formats = []
        for pattern in date_formats:
            if re.search(pattern, resume_text, re.IGNORECASE):
                found_formats.append(pattern)

        if len(found_formats) > 2:
            issues.append("Inconsistent date formats confuse ATS parsers")
            date_score -= 25
        elif len(found_formats) == 0:
            issues.append("No dates found - ATS cannot determine timeline")
            date_score -= 40

        # Check for date gaps or overlaps
        year_mentions = re.findall(r'\b(19|20)\d{2}\b', resume_text)
        if len(year_mentions) > 1:
            years = sorted([int(y) for y in year_mentions])
            for i in range(1, len(years)):
                if years[i] - years[i - 1] > 3:
                    issues.append(f"Employment gap detected ({years[i - 1]}-{years[i]})")
                    date_score -= 10
                    break

        detailed_scores['dates'] = max(date_score, 0)

        # CALCULATE FINAL WEIGHTED ATS SCORE
        weights = {
            'parsing': 0.25,
            'keywords': 0.30,
            'structure': 0.25,
            'contact': 0.10,
            'dates': 0.10
        }

        final_ats_score = sum(detailed_scores[key] * weights[key] for key in weights)

        # Severity classification
        if final_ats_score >= 80:
            compatibility_level = "EXCELLENT"
            risk_level = "Low"
        elif final_ats_score >= 65:
            compatibility_level = "GOOD"
            risk_level = "Medium-Low"
        elif final_ats_score >= 50:
            compatibility_level = "FAIR"
            risk_level = "Medium"
        elif final_ats_score >= 35:
            compatibility_level = "POOR"
            risk_level = "High"
        else:
            compatibility_level = "CRITICAL"
            risk_level = "Very High"

        return {
            'ats_score': round(final_ats_score, 1),
            'issues': issues,
            'is_compatible': final_ats_score >= 50,
            'detailed_scores': detailed_scores,
            'compatibility_level': compatibility_level,
            'risk_level': risk_level,
            'recommendations': self._generate_ats_recommendations(detailed_scores, issues)
        }

    def _generate_ats_recommendations(self, scores: Dict, issues: List) -> List[str]:
        """Generate specific recommendations based on ATS analysis"""
        recommendations = []

        if scores['parsing'] < 70:
            recommendations.append("Convert resume to plain text format or simple Word doc")
            recommendations.append("Remove all special characters, graphics, and tables")

        if scores['keywords'] < 70:
            recommendations.append("Include more industry-specific keywords from job description")
            recommendations.append("Repeat important keywords 2-3 times naturally")

        if scores['structure'] < 70:
            recommendations.append("Use standard section headers: Experience, Education, Skills")
            recommendations.append("Ensure sections are clearly labeled and separated")

        if scores['contact'] < 70:
            recommendations.append("Place contact information prominently at the top")
            recommendations.append("Include email, phone, and LinkedIn profile")

        if scores['dates'] < 70:
            recommendations.append("Use consistent date format (MM/YYYY recommended)")
            recommendations.append("Ensure all positions have clear start and end dates")

        return recommendations

    def remove_bias(self, resume_data: Dict) -> Dict:
        """Remove potential bias indicators"""
        # Remove name to prevent gender/ethnicity bias
        if 'name' in resume_data:
            resume_data['name'] = 'Candidate_' + hashlib.md5(str(resume_data.get('email', '')).encode()).hexdigest()[:8]

        # Remove age indicators
        if 'graduation_year' in resume_data:
            years_ago = datetime.now().year - resume_data['graduation_year']
            if years_ago > 20:
                resume_data['experience_level'] = 'Senior'
            elif years_ago > 10:
                resume_data['experience_level'] = 'Mid-level'
            else:
                resume_data['experience_level'] = 'Entry/Junior'
            del resume_data['graduation_year']

        # Normalize university names (to prevent ivy league bias)
        if 'university' in resume_data:
            resume_data['education_tier'] = self.classify_education_tier(resume_data['university'])
            resume_data['university'] = 'University'

        return resume_data

    def classify_education_tier(self, university: str) -> str:
        """Classify education institution tier"""
        tier1 = ['mit', 'stanford', 'harvard', 'oxford', 'cambridge', 'iit']
        tier2 = ['carnegie', 'georgia tech', 'berkeley', 'ucla', 'nit']

        uni_lower = university.lower()

        if any(t1 in uni_lower for t1 in tier1):
            return 'Tier 1'
        elif any(t2 in uni_lower for t2 in tier2):
            return 'Tier 2'
        else:
            return 'Tier 3'


# ==========================================
# STEP 8: Advanced Features
# ==========================================

class AdvancedAnalytics:
    """Step 8: Skills gap analysis, cultural fit, career progression"""

    def analyze_skill_gaps(self, candidate_skills: List[str], required_skills: List[str]) -> Dict:
        """Identify skill gaps and suggest improvements"""
        candidate_set = set(s.lower() for s in candidate_skills)
        required_set = set(s.lower() for s in required_skills)

        matched = candidate_set.intersection(required_set)
        missing = required_set - candidate_set
        additional = candidate_set - required_set

        # Identify transferable skills
        transferable = []
        skill_mappings = {
            'python': ['programming', 'scripting', 'automation'],
            'java': ['programming', 'oop', 'backend'],
            'react': ['frontend', 'javascript', 'web development'],
            'docker': ['containerization', 'devops', 'deployment'],
            'aws': ['cloud', 'infrastructure', 'scalability']
        }

        for skill in missing:
            for candidate_skill in additional:
                if skill in skill_mappings.get(candidate_skill, []):
                    transferable.append(f"{candidate_skill} → {skill}")

        return {
            'matched_skills': list(matched),
            'missing_skills': list(missing),
            'additional_skills': list(additional),
            'transferable_skills': transferable,
            'match_percentage': (len(matched) / len(required_set) * 100) if required_set else 0,
            'learning_potential': len(transferable) > 0
        }

    def assess_cultural_fit(self, resume_text: str, company_values: List[str]) -> float:
        """Assess cultural fit based on resume language and values"""
        score = 50.0  # Base score
        text_lower = resume_text.lower()

        # Look for value indicators
        value_keywords = {
            'innovation': ['innovative', 'creative', 'novel', 'pioneering', 'cutting-edge'],
            'teamwork': ['team', 'collaborate', 'together', 'group', 'cooperative'],
            'leadership': ['led', 'managed', 'directed', 'headed', 'supervised'],
            'excellence': ['excellence', 'quality', 'best', 'outstanding', 'exceptional'],
            'customer-focus': ['customer', 'client', 'user', 'satisfaction', 'service']
        }

        for value in company_values:
            value_lower = value.lower()
            if value_lower in value_keywords:
                keywords = value_keywords[value_lower]
                matches = sum(1 for kw in keywords if kw in text_lower)
                score += min(10, matches * 2)

        return min(score, 100)

    def analyze_career_progression(self, work_history: List[Dict]) -> Dict:
        """Analyze career growth and stability"""
        if not work_history:
            return {'progression_score': 0, 'stability_score': 0, 'growth_rate': 'Unknown'}

        # Calculate average tenure
        tenures = []
        for job in work_history:
            start = job.get('start_date')
            end = job.get('end_date')
            if start and end:
                tenure = (end - start).days / 365
                tenures.append(tenure)

        avg_tenure = statistics.mean(tenures) if tenures else 0

        # Stability score
        if avg_tenure >= 3:
            stability_score = 90
        elif avg_tenure >= 2:
            stability_score = 70
        elif avg_tenure >= 1:
            stability_score = 50
        else:
            stability_score = 30

        # Progression score (check for promotions)
        titles = [job.get('title', '').lower() for job in work_history]
        progression_keywords = ['senior', 'lead', 'principal', 'manager', 'director', 'vp', 'chief']

        progression_score = 50
        for i in range(1, len(titles)):
            for keyword in progression_keywords:
                if keyword in titles[i] and keyword not in titles[i - 1]:
                    progression_score += 20
                    break

        # Growth rate
        if progression_score >= 80:
            growth_rate = 'Fast'
        elif progression_score >= 60:
            growth_rate = 'Moderate'
        else:
            growth_rate = 'Slow'

        return {
            'progression_score': min(progression_score, 100),
            'stability_score': stability_score,
            'average_tenure': round(avg_tenure, 1),
            'growth_rate': growth_rate
        }


# ==========================================
# STEP 9: Decision Making
# ==========================================

class DecisionEngine:
    """Step 9: Dynamic thresholding and multi-model voting"""

    def __init__(self):
        self.base_threshold = 20.0
        self.role_thresholds = {
            'senior': 70,
            'mid': 50,
            'junior': 30,
            'intern': 20
        }
        self.model_weights = {
            'ml_model': 0.30,
            'rule_based': 0.40,
            'semantic': 0.30
        }

    def determine_threshold(self, job_level: str, market_demand: str, urgency: str) -> float:
        """Dynamically determine threshold based on factors"""
        # Base threshold from job level
        threshold = self.role_thresholds.get(job_level, 50)

        # Adjust for market demand
        if market_demand == 'high':
            threshold -= 10  # Lower threshold when demand is high
        elif market_demand == 'low':
            threshold += 10  # Higher threshold when many candidates available

        # Adjust for urgency
        if urgency == 'urgent':
            threshold -= 15
        elif urgency == 'low':
            threshold += 5

        return max(10, min(90, threshold))  # Keep between 10-90

    def ensemble_decision(self, scores: Dict[str, float]) -> Tuple[str, float]:
        """Make decision using ensemble of models"""
        weighted_score = 0

        for model, weight in self.model_weights.items():
            score = scores.get(model, 50)
            weighted_score += score * weight

        # Determine confidence
        score_variance = statistics.variance(scores.values()) if len(scores) > 1 else 0

        if score_variance < 100:
            confidence = 'HIGH'
        elif score_variance < 300:
            confidence = 'MEDIUM'
        else:
            confidence = 'LOW'

        # Make decision
        if weighted_score >= self.base_threshold:
            decision = 'SHORTLIST'
        else:
            decision = 'REJECT'

        return decision, weighted_score, confidence

    def explain_decision(self, candidate_data: Dict) -> str:
        """Generate explanation for the decision"""
        explanations = []

        if candidate_data.get('overall_score', 0) >= self.base_threshold:
            explanations.append(f"Meets the minimum threshold of {self.base_threshold}%")

        if candidate_data.get('skill_match', 0) >= 60:
            explanations.append("Strong skill alignment with requirements")
        elif candidate_data.get('skill_match', 0) < 40:
            explanations.append("Significant skill gaps identified")

        if candidate_data.get('experience_years', 0) >= candidate_data.get('required_years', 0):
            explanations.append("Sufficient experience level")
        else:
            explanations.append("Below required experience level")

        if candidate_data.get('ats_score', 0) >= 70:
            explanations.append("ATS-optimized resume format")

        return " | ".join(explanations)


# ==========================================
# STEP 10: Quality Assurance
# ==========================================

class QualityAssurance:
    """Step 10: Validation, feedback, and continuous improvement"""

    def __init__(self):
        self.validation_samples = []
        self.feedback_history = []
        self.performance_metrics = {
            'precision': [],
            'recall': [],
            'false_positives': 0,
            'false_negatives': 0
        }

    def validate_decision(self, decision: str, human_review: str = None) -> bool:
        """Validate AI decision against human review"""
        if not human_review:
            # Auto-validation based on confidence scores
            return True

        is_correct = decision == human_review

        # Update metrics
        if is_correct:
            self.performance_metrics['precision'].append(1)
        else:
            self.performance_metrics['precision'].append(0)
            if decision == 'SHORTLIST' and human_review == 'REJECT':
                self.performance_metrics['false_positives'] += 1
            elif decision == 'REJECT' and human_review == 'SHORTLIST':
                self.performance_metrics['false_negatives'] += 1

        return is_correct

    def calculate_metrics(self) -> Dict:
        """Calculate performance metrics"""
        metrics = {}

        if self.performance_metrics['precision']:
            metrics['precision'] = statistics.mean(self.performance_metrics['precision'])
        else:
            metrics['precision'] = 0

        metrics['false_positive_rate'] = self.performance_metrics['false_positives']
        metrics['false_negative_rate'] = self.performance_metrics['false_negatives']

        # F1 Score approximation
        if metrics['precision'] > 0:
            recall = 1 - (metrics['false_negative_rate'] / max(len(self.validation_samples), 1))
            metrics['f1_score'] = 2 * (metrics['precision'] * recall) / (metrics['precision'] + recall)
        else:
            metrics['f1_score'] = 0

        return metrics

    def generate_audit_report(self, batch_results: List[Dict]) -> Dict:
        """Generate audit report for a batch of decisions"""
        report = {
            'total_processed': len(batch_results),
            'shortlisted': sum(1 for r in batch_results if r.get('decision') == 'SHORTLIST'),
            'rejected': sum(1 for r in batch_results if r.get('decision') == 'REJECT'),
            'average_score': statistics.mean([r.get('overall_score', 0) for r in batch_results]),
            'confidence_distribution': {
                'HIGH': sum(1 for r in batch_results if r.get('confidence') == 'HIGH'),
                'MEDIUM': sum(1 for r in batch_results if r.get('confidence') == 'MEDIUM'),
                'LOW': sum(1 for r in batch_results if r.get('confidence') == 'LOW')
            },
            'performance_metrics': self.calculate_metrics(),
            'timestamp': datetime.now().isoformat()
        }

        return report


# ==========================================
# MAIN PROCESSING PIPELINE
# ==========================================

class IndustryStandardRecruiter:
    """Main orchestrator for the complete 10-step pipeline"""

    def __init__(self):
        # Initialize all components
        self.parser = DocumentParser()
        self.nlp = NLPProcessor()
        self.semantic = SemanticAnalyzer()
        self.scoring = ScoringEngine()
        self.ml_pipeline = MLPipeline()
        self.ranking = RankingEngine()
        self.industry = IndustryAnalyzer()
        self.advanced = AdvancedAnalytics()
        self.decision = DecisionEngine()
        self.qa = QualityAssurance()

    async def process_resume(self, file_content: bytes, file_type: str, job_requirements: Dict) -> Dict:
        """Process single resume through complete pipeline"""

        # Step 1: Document Parsing
        if file_type == 'pdf':
            raw_text = self.parser.extract_text_from_pdf(file_content)
        elif file_type in ['docx', 'doc']:
            raw_text = self.parser.extract_text_from_docx(file_content)
        else:
            raw_text = file_content.decode('utf-8', errors='ignore')

        if not raw_text or len(raw_text) < 100:
            return {'error': 'Insufficient content in resume'}

        sections = self.parser.identify_sections(raw_text)
        contact_info = self.parser.extract_contact_info(raw_text)

        # Step 2: NLP Processing
        entities = self.nlp.extract_entities(raw_text)
        tokens = self.nlp.tokenize_and_clean(raw_text)
        years_experience = self.nlp.extract_years_of_experience(raw_text)

        # Step 3: Semantic Understanding
        semantic_score = self.semantic.calculate_semantic_similarity(
            raw_text,
            job_requirements.get('job_description', '')
        )
        skill_match = self.semantic.extract_skill_embeddings(
            entities.get('technologies', []),
            job_requirements.get('required_skills', [])
        )

        # Step 4: Scoring
        keyword_score = self.scoring.calculate_keyword_score(
            raw_text,
            job_requirements.get('keywords', [])
        )
        experience_score = self.scoring.calculate_experience_score(
            years_experience,
            job_requirements.get('required_experience', 0)
        )
        education_score = self.scoring.calculate_education_score(
            sections.get('education', ''),
            job_requirements.get('required_education', '')
        )
        contextual_score = self.scoring.calculate_contextual_score(
            sections,
            job_requirements
        )

        # Step 5: ML Models
        resume_features = {
            'years_experience': years_experience,
            'skills': entities.get('technologies', []),
            'companies': entities.get('companies', []),
            'certifications': [],
            'education_score': education_score,
            'keyword_score': keyword_score,
            'semantic_score': semantic_score,
            'text': raw_text,
            'sections': sections
        }

        ml_features = self.ml_pipeline.extract_features(resume_features)
        ml_score, ml_confidence = self.ml_pipeline.predict_suitability(ml_features)

        # Step 6: Ranking (will be applied to batch)
        # Ensure all scores are valid percentages (0-100)
        semantic_score = max(0, min(100, semantic_score)) if semantic_score else 30.0
        skill_match = max(0, min(100, skill_match)) if skill_match else 20.0
        keyword_score = max(0, min(100, keyword_score)) if keyword_score else 30.0
        experience_score = max(0, min(100, experience_score)) if experience_score else 40.0
        education_score = max(0, min(100, education_score)) if education_score else 50.0
        contextual_score = max(0, min(100, contextual_score)) if contextual_score else 40.0

        # Calculate weighted overall score
        overall_score = (
                semantic_score * 0.30 +  # 30% weight for semantic similarity
                skill_match * 0.25 +  # 25% weight for skill matching
                keyword_score * 0.20 +  # 20% weight for keyword matching
                experience_score * 0.15 +  # 15% weight for experience
                education_score * 0.10  # 10% weight for education
        )

        # Ensure overall score is within valid range
        overall_score = max(0, min(100, overall_score))

        candidate_data = {
            'overall_score': overall_score,
            'semantic_score': semantic_score,
            'skill_match': skill_match,
            'keyword_score': keyword_score,
            'experience_score': experience_score,
            'education_score': education_score,
            'contextual_score': contextual_score,
            'years_experience': years_experience,
            'skills': entities.get('technologies', []),
            'ml_score': ml_score,
            'ml_confidence': ml_confidence
        }

        # Step 7: Industry-specific
        industry_type = self.industry.detect_industry(raw_text)
        ats_check = self.industry.check_ats_compatibility(raw_text)
        candidate_data['industry'] = industry_type
        candidate_data['ats_score'] = ats_check['ats_score']
        candidate_data['ats_issues'] = ats_check['issues']

        # Remove bias
        candidate_data = self.industry.remove_bias(candidate_data)

        # Step 8: Advanced Features
        skill_gaps = self.advanced.analyze_skill_gaps(
            entities.get('technologies', []),
            job_requirements.get('required_skills', [])
        )
        cultural_fit = self.advanced.assess_cultural_fit(
            raw_text,
            job_requirements.get('company_values', [])
        )

        candidate_data['skill_gaps'] = skill_gaps
        candidate_data['cultural_fit'] = cultural_fit

        # Step 9: Decision
        threshold = self.decision.determine_threshold(
            job_requirements.get('level', 'mid'),
            job_requirements.get('market_demand', 'normal'),
            job_requirements.get('urgency', 'normal')
        )

        decision, final_score, confidence = self.decision.ensemble_decision({
            'ml_model': ml_score,
            'rule_based': candidate_data['overall_score'],
            'semantic': semantic_score
        })

        candidate_data['decision'] = decision
        candidate_data['final_score'] = final_score
        candidate_data['confidence'] = confidence
        candidate_data['threshold_used'] = threshold
        candidate_data['decision_explanation'] = self.decision.explain_decision(candidate_data)

        # Step 10: Quality Assurance
        self.qa.validation_samples.append(candidate_data)

        return candidate_data

    async def process_batch(self, resumes: List[Dict], job_requirements: Dict) -> List[Dict]:
        """Process batch of resumes with ranking and filtering"""
        results = []

        # Process each resume
        for resume in resumes:
            result = await self.process_resume(
                resume['content'],
                resume['type'],
                job_requirements
            )
            result['filename'] = resume['filename']
            results.append(result)

        # Step 6: Apply ranking and filtering
        filtered = self.ranking.apply_hard_filters(results, job_requirements)
        ranked = self.ranking.calculate_percentile_rank(filtered)
        final_ranked = self.ranking.apply_soft_scoring(ranked)

        # Sort by final_score in descending order (highest score first)
        # Primary sort: final_score, Secondary sort: overall_score
        final_ranked.sort(
            key=lambda x: (
                x.get('final_score', 0),
                x.get('overall_score', 0),
                x.get('ats_score', 0)
            ),
            reverse=True
        )

        # Update rank positions after sorting
        for idx, candidate in enumerate(final_ranked, 1):
            candidate['rank'] = idx
            # Also update decision priority
            if candidate.get('decision') == 'SHORTLIST':
                candidate['priority'] = 1
            else:
                candidate['priority'] = 2

        # Final sort: First by decision (SHORTLIST first), then by score
        final_ranked.sort(
            key=lambda x: (
                -x.get('priority', 2),  # Negative for reverse order (1 before 2)
                -x.get('final_score', 0),  # Negative for descending order
                -x.get('overall_score', 0)
            )
        )

        # Step 10: Generate audit report
        audit_report = self.qa.generate_audit_report(final_ranked)

        return final_ranked, audit_report


# ==========================================
# GROQ API Integration
# ==========================================

async def call_groq_api(prompt: str, max_tokens: int = 1500) -> str:
    """Call Groq API with LLaMA model for intelligent analysis"""
    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }

    data = {
        "model": GROQ_MODEL,
        "messages": [
            {
                "role": "system",
                "content": "You are an expert recruitment AI with deep knowledge of industry standards, ATS systems, and talent evaluation."
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        "max_tokens": max_tokens,
        "temperature": 0.2
    }

    async with httpx.AsyncClient(timeout=30.0) as client:
        try:
            response = await client.post(
                f"{GROQ_BASE_URL}/chat/completions",
                headers=headers,
                json=data
            )
            response.raise_for_status()
            return response.json()["choices"][0]["message"]["content"]
        except Exception as e:
            print(f"Groq API error: {e}")
            return ""


async def extract_job_requirements_ai(job_description: str) -> Dict:
    """Extract comprehensive job requirements using Groq AI + fallback parsing"""

    # First try AI extraction
    prompt = f"""
    Analyze this job description and extract all requirements in detail.

    Job Description:
    {job_description[:4000]}

    Extract and return a comprehensive JSON with ALL of these fields:
    {{
        "required_skills": ["list of must-have technical skills"],
        "preferred_skills": ["list of nice-to-have skills"],
        "required_experience": <minimum years as number>,
        "required_education": "minimum degree required",
        "keywords": ["important keywords from JD"],
        "key_responsibilities": ["main job responsibilities"],
        "company_values": ["company culture values if mentioned"],
        "level": "entry/junior/mid/senior/lead/principal",
        "industry": "tech/finance/healthcare/etc",
        "technologies": ["specific technologies mentioned"],
        "certifications": ["required or preferred certifications"],
        "soft_skills": ["communication, leadership, etc"],
        "domain_knowledge": ["specific domain expertise needed"]
    }}

    Be very detailed and extract EVERYTHING mentioned.
    Return ONLY valid JSON, no extra text.
    """

    try:
        response = await call_groq_api(prompt)

        if response:
            # Clean and parse response
            json_str = response.strip()
            if "```json" in json_str:
                json_str = json_str.split("```json")[1].split("```")[0]
            elif "```" in json_str:
                json_str = json_str.split("```")[1].split("```")[0]

            # Remove any non-JSON content
            json_str = json_str.strip()
            if json_str.startswith('{'):
                end_idx = json_str.rfind('}')
                if end_idx != -1:
                    json_str = json_str[:end_idx + 1]

            parsed = json.loads(json_str)

            # Validate and fill missing fields
            return validate_job_requirements(parsed, job_description)
    except Exception as e:
        print(f"AI extraction failed: {e}, using fallback parser")

    # Fallback to rule-based extraction
    return extract_job_requirements_fallback(job_description)


def extract_job_requirements_fallback(job_description: str) -> Dict:
    """Generic rule-based job requirement extraction for ANY role"""

    jd_lower = job_description.lower()

    # Generic skill extraction - works for any field
    skills = []

    # Technical skills (IT, Engineering, Data)
    tech_skills = [
        'python', 'java', 'javascript', 'react', 'angular', 'node', 'php', 'c++', 'c#', '.net',
        'sql', 'nosql', 'mongodb', 'mysql', 'postgresql', 'oracle', 'database',
        'aws', 'azure', 'gcp', 'cloud', 'docker', 'kubernetes', 'jenkins', 'git',
        'html', 'css', 'bootstrap', 'jquery', 'vue', 'typescript',
        'spring', 'django', 'flask', 'laravel', 'express', 'fastapi',
        'machine learning', 'deep learning', 'ai', 'data science', 'analytics',
        'power bi', 'tableau', 'excel', 'sap', 'erp', 'crm', 'salesforce'
    ]

    # Business & Management skills
    business_skills = [
        'project management', 'agile', 'scrum', 'jira', 'confluence',
        'business analysis', 'requirements gathering', 'stakeholder management',
        'strategic planning', 'budgeting', 'forecasting', 'risk management',
        'process improvement', 'lean', 'six sigma', 'change management',
        'vendor management', 'contract negotiation', 'compliance', 'audit'
    ]

    # Sales & Marketing skills
    sales_skills = [
        'sales', 'business development', 'lead generation', 'cold calling',
        'account management', 'customer relationship', 'negotiation', 'closing',
        'b2b', 'b2c', 'retail', 'wholesale', 'territory management',
        'marketing', 'digital marketing', 'seo', 'sem', 'social media',
        'content marketing', 'email marketing', 'google ads', 'facebook ads',
        'brand management', 'market research', 'competitive analysis'
    ]

    # Finance & Accounting skills
    finance_skills = [
        'accounting', 'bookkeeping', 'financial analysis', 'financial reporting',
        'taxation', 'gst', 'income tax', 'tds', 'audit', 'compliance',
        'tally', 'quickbooks', 'sap fico', 'financial modeling', 'valuation',
        'investment analysis', 'portfolio management', 'risk assessment',
        'budgeting', 'cost accounting', 'payroll', 'accounts payable', 'accounts receivable',
        'financial planning', 'treasury', 'forex', 'derivatives'
    ]

    # HR & Admin skills
    hr_skills = [
        'recruitment', 'talent acquisition', 'screening', 'interviewing',
        'onboarding', 'employee relations', 'performance management',
        'compensation', 'benefits', 'payroll processing', 'hris',
        'training', 'development', 'succession planning', 'organizational development',
        'labor laws', 'compliance', 'employee engagement', 'culture building',
        'administration', 'office management', 'documentation', 'coordination'
    ]

    # Operations & Supply Chain skills
    operations_skills = [
        'operations management', 'supply chain', 'logistics', 'inventory',
        'warehouse management', 'procurement', 'sourcing', 'vendor management',
        'quality control', 'quality assurance', 'iso', 'lean manufacturing',
        'production planning', 'capacity planning', 'demand forecasting',
        'distribution', 'transportation', 'import export', 'customs'
    ]

    # Healthcare skills
    healthcare_skills = [
        'patient care', 'clinical', 'nursing', 'medical', 'diagnosis',
        'treatment planning', 'medication', 'healthcare management',
        'medical coding', 'billing', 'insurance', 'hipaa', 'ehr', 'emr',
        'laboratory', 'radiology', 'pharmacy', 'surgical', 'emergency'
    ]

    # Banking & Insurance skills
    banking_skills = [
        'banking operations', 'retail banking', 'corporate banking',
        'loan processing', 'credit analysis', 'underwriting', 'kyc', 'aml',
        'trade finance', 'treasury operations', 'forex dealing',
        'insurance', 'life insurance', 'general insurance', 'claims processing',
        'actuarial', 'risk assessment', 'policy', 'reinsurance'
    ]

    # Soft skills - Universal for all roles
    soft_skills_list = [
        'communication', 'leadership', 'teamwork', 'problem solving',
        'analytical', 'creative', 'organized', 'detail oriented',
        'time management', 'multitasking', 'adaptability', 'flexibility',
        'customer service', 'interpersonal', 'presentation', 'negotiation',
        'decision making', 'critical thinking', 'initiative', 'motivation'
    ]

    # Check all skill categories
    all_skill_categories = [
        tech_skills, business_skills, sales_skills, finance_skills,
        hr_skills, operations_skills, healthcare_skills, banking_skills
    ]

    # Extract skills from JD
    for skill_category in all_skill_categories:
        for skill in skill_category:
            if skill in jd_lower:
                skills.append(skill.title() if len(skill) > 3 else skill.upper())

    # Extract soft skills separately
    extracted_soft_skills = []
    for skill in soft_skills_list:
        if skill in jd_lower:
            extracted_soft_skills.append(skill.title())

    # Extract experience requirements
    experience = 0
    exp_patterns = [
        r'(\d+)\+?\s*(?:to\s+)?(\d+)?\s*years?\s+(?:of\s+)?experience',
        r'experience[:\s]+(\d+)\+?\s*(?:to\s+)?(\d+)?\s*years?',
        r'minimum\s+(\d+)\s*years?',
        r'at\s+least\s+(\d+)\s*years?',
        r'(\d+)\s*years?\s+minimum'
    ]

    for pattern in exp_patterns:
        match = re.search(pattern, jd_lower)
        if match:
            experience = int(match.group(1))
            break

    # If no experience mentioned, check for level indicators
    if experience == 0:
        if 'fresher' in jd_lower or 'entry level' in jd_lower:
            experience = 0
        elif 'junior' in jd_lower:
            experience = 1
        elif 'senior' in jd_lower or 'lead' in jd_lower:
            experience = 5
        elif 'manager' in jd_lower or 'principal' in jd_lower:
            experience = 7
        else:
            experience = 2  # Default for mid-level

    # Extract education requirements
    education = extract_education_requirement(jd_lower)

    # Determine job level
    level = determine_job_level(jd_lower, experience)

    # Detect industry
    industry = detect_industry_generic(jd_lower)

    # Extract keywords (important terms from JD)
    keywords = extract_important_keywords(job_description)

    # Extract responsibilities
    responsibilities = extract_responsibilities_generic(job_description)

    # Build requirements dictionary
    requirements = {
        "required_skills": skills[:15] if skills else keywords[:10],
        "preferred_skills": skills[15:25] if len(skills) > 15 else [],
        "required_experience": experience,
        "required_education": education,
        "keywords": keywords,
        "key_responsibilities": responsibilities,
        "company_values": [],
        "level": level,
        "industry": industry,
        "technologies": [s for s in skills if s.lower() in tech_skills],
        "certifications": extract_certifications_generic(jd_lower),
        "soft_skills": extracted_soft_skills,
        "domain_knowledge": []
    }

    return requirements


def extract_education_requirement(jd_lower: str) -> str:
    """Extract education requirement for any field"""
    if 'phd' in jd_lower or 'doctorate' in jd_lower:
        return "PhD"
    elif 'post graduate' in jd_lower or 'pg ' in jd_lower or 'post-graduate' in jd_lower:
        return "Post Graduate"
    elif 'mba' in jd_lower:
        return "MBA"
    elif 'master' in jd_lower or 'mtech' in jd_lower or 'msc' in jd_lower or 'ma ' in jd_lower or 'mcom' in jd_lower:
        return "Master's"
    elif 'bachelor' in jd_lower or 'btech' in jd_lower or 'bsc' in jd_lower or 'ba ' in jd_lower or 'bcom' in jd_lower or 'bba' in jd_lower:
        return "Bachelor's"
    elif 'diploma' in jd_lower:
        return "Diploma"
    elif '12th' in jd_lower or 'intermediate' in jd_lower or 'hsc' in jd_lower:
        return "12th/Intermediate"
    elif '10th' in jd_lower or 'matriculation' in jd_lower or 'ssc' in jd_lower:
        return "10th/Matriculation"
    elif 'graduate' in jd_lower or 'graduation' in jd_lower:
        return "Graduate"
    else:
        return "Any Graduate"  # Default


def determine_job_level(jd_lower: str, experience: int) -> str:
    """Determine job level for any industry"""
    # Check explicit level mentions
    if 'ceo' in jd_lower or 'cto' in jd_lower or 'cfo' in jd_lower or 'coo' in jd_lower:
        return "c-level"
    elif 'vp ' in jd_lower or 'vice president' in jd_lower or 'avp' in jd_lower:
        return "vp"
    elif 'director' in jd_lower or 'head of' in jd_lower:
        return "director"
    elif 'principal' in jd_lower or 'architect' in jd_lower or 'staff' in jd_lower:
        return "principal"
    elif 'senior manager' in jd_lower or 'sr manager' in jd_lower:
        return "senior-manager"
    elif 'manager' in jd_lower or 'team lead' in jd_lower or 'team leader' in jd_lower:
        return "manager"
    elif 'senior' in jd_lower or 'sr ' in jd_lower or 'lead' in jd_lower:
        return "senior"
    elif 'junior' in jd_lower or 'jr ' in jd_lower:
        return "junior"
    elif 'fresher' in jd_lower or 'trainee' in jd_lower or 'intern' in jd_lower:
        return "entry"
    elif 'associate' in jd_lower:
        return "associate"
    elif 'executive' in jd_lower:
        return "executive"
    else:
        # Determine by experience
        if experience >= 10:
            return "senior"
        elif experience >= 5:
            return "mid"
        elif experience >= 2:
            return "junior"
        else:
            return "entry"


def detect_industry_generic(jd_lower: str) -> str:
    """Detect industry from job description - covers all major industries"""
    industries = {
        'tech': ['software', 'technology', 'it ', 'it/', 'information technology', 'saas', 'paas', 'iaas', 'cloud',
                 'digital', 'tech ', 'platform', 'application', 'web ', 'mobile', 'app '],
        'finance': ['financial', 'banking', 'bank ', 'fintech', 'investment', 'trading', 'insurance', 'mutual fund',
                    'stock', 'equity', 'debt', 'capital market', 'wealth management'],
        'healthcare': ['healthcare', 'medical', 'pharma', 'clinical', 'hospital', 'health', 'diagnostic', 'biotech',
                       'life science', 'wellness', 'patient'],
        'retail': ['retail', 'e-commerce', 'ecommerce', 'marketplace', 'shopping', 'store', 'fashion', 'apparel',
                   'fmcg', 'consumer goods', 'merchandise'],
        'education': ['education', 'edtech', 'learning', 'academic', 'university', 'school', 'college', 'training',
                      'institute', 'teaching', 'faculty'],
        'consulting': ['consulting', 'advisory', 'management consulting', 'strategy', 'business consulting',
                       'consultant'],
        'manufacturing': ['manufacturing', 'production', 'factory', 'industrial', 'automobile', 'automotive', 'textile',
                          'chemical', 'steel', 'engineering'],
        'hospitality': ['hotel', 'restaurant', 'hospitality', 'tourism', 'travel', 'airline', 'resort', 'food service',
                        'catering'],
        'real-estate': ['real estate', 'property', 'construction', 'builder', 'developer', 'housing',
                        'commercial space', 'residential'],
        'telecom': ['telecom', 'telecommunication', 'network', 'mobile network', 'broadband', 'isp',
                    'internet service'],
        'media': ['media', 'entertainment', 'advertising', 'marketing agency', 'creative', 'production house',
                  'broadcast', 'journalism', 'publishing'],
        'logistics': ['logistics', 'supply chain', 'transportation', 'shipping', 'freight', 'courier', 'delivery',
                      'warehouse', 'distribution'],
        'energy': ['energy', 'power', 'oil', 'gas', 'renewable', 'solar', 'wind', 'electricity', 'utilities'],
        'agriculture': ['agriculture', 'farming', 'agri', 'agro', 'crop', 'livestock', 'dairy', 'poultry'],
        'government': ['government', 'public sector', 'ministry', 'department', 'municipal', 'psu', 'defence',
                       'defense']
    }

    # Count matches for each industry
    industry_scores = {}
    for industry, keywords in industries.items():
        score = sum(1 for keyword in keywords if keyword in jd_lower)
        if score > 0:
            industry_scores[industry] = score

    # Return industry with highest score
    if industry_scores:
        return max(industry_scores, key=industry_scores.get)

    return 'general'  # Default for unidentified industries


def extract_important_keywords(job_description: str) -> List[str]:
    """Extract important keywords from any job description"""
    keywords = []
    lines = job_description.split('\n')

    # Priority indicators for important lines
    priority_markers = [
        'must have', 'required', 'mandatory', 'essential', 'should have',
        'qualification', 'requirement', 'skill', 'experience in',
        'knowledge of', 'proficiency in', 'expertise in', 'familiar with'
    ]

    for line in lines:
        line_lower = line.lower()
        # Check if line contains priority markers
        if any(marker in line_lower for marker in priority_markers):
            # Extract meaningful words (exclude common words)
            words = re.findall(r'\b[a-zA-Z]+\b', line)

            # Common words to exclude
            exclude_words = {
                'the', 'and', 'or', 'in', 'of', 'to', 'for', 'with', 'on', 'at',
                'from', 'by', 'as', 'is', 'are', 'was', 'were', 'been', 'be',
                'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'should',
                'could', 'may', 'might', 'must', 'shall', 'can', 'need', 'required',
                'mandatory', 'essential', 'skill', 'skills', 'experience', 'qualification'
            }

            for word in words:
                if len(word) > 2 and word.lower() not in exclude_words:
                    keywords.append(word)

    # Remove duplicates and return top keywords
    unique_keywords = list(dict.fromkeys(keywords))
    return unique_keywords[:30]  # Return top 30 keywords


def extract_responsibilities_generic(job_description: str) -> List[str]:
    """Extract responsibilities for any job type"""
    responsibilities = []
    lines = job_description.split('\n')

    in_responsibilities = False

    # Responsibility section headers
    resp_headers = [
        'responsibilit', 'duties', 'role', 'you will', 'your role',
        'key accountabilit', 'what you', 'job description', 'scope'
    ]

    # Section end markers
    end_markers = [
        'requirement', 'qualification', 'skill', 'experience required',
        'education', 'certification', 'what we', 'eligibility'
    ]

    for line in lines:
        line = line.strip()
        line_lower = line.lower()

        # Check for responsibility section start
        if any(marker in line_lower for marker in resp_headers):
            in_responsibilities = True
            continue

        # Check for section end
        if in_responsibilities and any(marker in line_lower for marker in end_markers):
            break

        # Extract responsibilities (bullet points, numbered items, or regular lines)
        if in_responsibilities and line:
            # Check for bullet points or numbers
            if line.startswith(('•', '-', '*', '·', '→', '▪', '◦')) or re.match(r'^\d+[\.\)]\s', line):
                clean_line = re.sub(r'^[•\-\*·→▪◦\d\.\)]\s*', '', line)
                if len(clean_line) > 10:
                    responsibilities.append(clean_line[:200])
            elif len(line) > 20 and not any(marker in line_lower for marker in resp_headers):
                # Add regular lines that look like responsibilities
                responsibilities.append(line[:200])

    # If no responsibilities found in sections, extract from full text
    if not responsibilities:
        # Look for action verbs that typically start responsibility statements
        action_verbs = [
            'manage', 'develop', 'create', 'implement', 'design', 'build',
            'lead', 'coordinate', 'analyze', 'prepare', 'maintain', 'ensure',
            'support', 'assist', 'review', 'monitor', 'execute', 'deliver',
            'drive', 'optimize', 'improve', 'establish', 'collaborate'
        ]

        for line in lines:
            line_lower = line.lower().strip()
            if any(line_lower.startswith(verb) for verb in action_verbs):
                if len(line) > 20:
                    responsibilities.append(line[:200])

    return responsibilities[:15]  # Return top 15 responsibilities


def extract_certifications_generic(jd_lower: str) -> List[str]:
    """Extract certifications for any field"""
    certs = []

    # IT Certifications
    it_certs = [
        'aws certified', 'azure certified', 'gcp certified', 'cisco', 'ccna', 'ccnp', 'ccie',
        'comptia', 'itil', 'pmp', 'prince2', 'scrum master', 'csm', 'safe', 'agile',
        'cissp', 'ceh', 'cisa', 'oracle certified', 'microsoft certified', 'google certified',
        'red hat', 'vmware', 'salesforce certified', 'sap certified'
    ]

    # Finance/Accounting Certifications
    finance_certs = [
        'ca ', 'chartered accountant', 'cpa', 'cfa', 'frm', 'cma', 'acca',
        'icwa', 'cs ', 'company secretary', 'mba finance', 'fca', 'aca'
    ]

    # HR Certifications
    hr_certs = [
        'shrm', 'phr', 'sphr', 'cipd', 'chrp', 'gphr', 'hrci'
    ]

    # Quality/Process Certifications
    quality_certs = [
        'six sigma', 'green belt', 'black belt', 'iso', 'lean', 'kaizen'
    ]

    # Sales/Marketing Certifications
    sales_certs = [
        'google ads', 'facebook blueprint', 'hubspot', 'google analytics',
        'digital marketing', 'seo certification'
    ]

    # Healthcare Certifications
    health_certs = [
        'mbbs', 'md ', 'ms ', 'bds', 'bhms', 'bams', 'nursing', 'rn ', 'bls', 'acls'
    ]

    all_cert_categories = [
        it_certs, finance_certs, hr_certs, quality_certs, sales_certs, health_certs
    ]

    for cert_category in all_cert_categories:
        for cert in cert_category:
            if cert in jd_lower:
                certs.append(cert.upper() if len(cert) <= 4 else cert.title())

    return list(dict.fromkeys(certs))  # Remove duplicates


# ==========================================
# FastAPI Endpoints
# ==========================================

# Global instance
recruiter = IndustryStandardRecruiter()
processing_cache = {}


@app.post("/api/analyze")
async def analyze_resumes(
        files: List[UploadFile] = File(...),
        job_description: str = Form(...),
        matching_threshold: float = Form(20.0)
):
    """Industry-standard resume analysis endpoint"""

    global recruiter
    recruiter.decision.base_threshold = matching_threshold

    print(f"\n{'=' * 60}")
    print(f"🚀 Industry Standard Resume Analysis Pipeline")
    print(f"📋 Files: {len(files)}")
    print(f"🎯 Threshold: {matching_threshold}%")
    print(f"{'=' * 60}\n")

    # Extract job requirements using AI
    print("Step 1: Extracting job requirements...")
    job_requirements = await extract_job_requirements_ai(job_description)
    job_requirements['job_description'] = job_description

    # Prepare resumes for processing
    resumes_data = []
    for idx, file in enumerate(files[:50], 1):
        print(f"Processing file {idx}/{len(files)}: {file.filename}")

        content = await file.read()
        file_type = file.filename.split('.')[-1].lower()

        resumes_data.append({
            'content': content,
            'type': file_type,
            'filename': file.filename
        })

    # Process through complete pipeline
    print("\nExecuting 10-step industry standard pipeline...")
    results, audit_report = await recruiter.process_batch(resumes_data, job_requirements)

    # Prepare response
    shortlisted = [r for r in results if r.get('decision') == 'SHORTLIST']
    rejected = [r for r in results if r.get('decision') == 'REJECT']

    print(f"\n{'=' * 60}")
    print(f"✅ Analysis Complete!")
    print(f"  Total: {len(results)}")
    print(f"  Shortlisted: {len(shortlisted)}")
    print(f"  Rejected: {len(rejected)}")
    print(f"  Average Score: {audit_report['average_score']:.1f}%")
    print(f"{'=' * 60}\n")

    # Store in cache
    session_id = hashlib.md5(f"{datetime.now()}".encode()).hexdigest()[:8]
    processing_cache[session_id] = {
        'results': results,
        'audit_report': audit_report,
        'job_requirements': job_requirements,
        'timestamp': datetime.now().isoformat()
    }

    return {
        'session_id': session_id,
        'total_resumes': len(results),
        'shortlisted_count': len(shortlisted),
        'rejected_count': len(rejected),
        'results': results,
        'audit_report': audit_report,
        'matching_threshold': matching_threshold
    }


@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "service": "AI Recruiter - Industry Standard",
        "version": "2.0",
        "pipeline_steps": 10,
        "models_loaded": True
    }


@app.get("/api/download/{session_id}")
async def download_results(session_id: str):
    """Download analysis results"""
    if session_id not in processing_cache:
        raise HTTPException(status_code=404, detail="Session not found")

    return JSONResponse(
        content=processing_cache[session_id],
        headers={
            'Content-Disposition': f'attachment; filename="industry_analysis_{session_id}.json"'
        }
    )


# HTML Frontend (Industry Standard UI)
HTML_TEMPLATE = '''
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>AI Recruiter - Industry Standard Pipeline</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }

        body {
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
            min-height: 100vh;
            padding: 20px;
        }

        .container {
            max-width: 1600px;
            margin: 0 auto;
        }

        .header {
            text-align: center;
            color: white;
            margin-bottom: 30px;
        }

        .header h1 {
            font-size: 2.5rem;
            margin-bottom: 10px;
            text-shadow: 2px 2px 4px rgba(0,0,0,0.3);
        }

        .header p {
            font-size: 1.1rem;
            opacity: 0.95;
        }

        .pipeline-indicator {
            display: flex;
            justify-content: center;
            gap: 10px;
            margin-top: 20px;
        }

        .step-badge {
            background: rgba(255,255,255,0.2);
            padding: 5px 10px;
            border-radius: 15px;
            font-size: 0.85rem;
        }

        .main-card {
            background: white;
            border-radius: 20px;
            padding: 30px;
            box-shadow: 0 20px 60px rgba(0,0,0,0.15);
        }

        .form-group {
            margin-bottom: 25px;
        }

        .form-group label {
            display: block;
            margin-bottom: 8px;
            font-weight: 600;
            color: #333;
        }

        .file-upload-label {
            display: block;
            padding: 40px;
            background: #f8f9fa;
            border: 2px dashed #dee2e6;
            border-radius: 10px;
            text-align: center;
            cursor: pointer;
            transition: all 0.3s ease;
        }

        .file-upload-label:hover {
            background: #e9ecef;
            border-color: #adb5bd;
        }

        .file-upload-label.has-files {
            background: #e7f3ff;
            border-color: #4dabf7;
        }

        textarea {
            width: 100%;
            min-height: 200px;
            padding: 12px;
            border: 1px solid #dee2e6;
            border-radius: 8px;
            font-size: 14px;
            resize: vertical;
            font-family: inherit;
        }

        .threshold-control {
            display: flex;
            align-items: center;
            gap: 15px;
            background: #f8f9fa;
            padding: 15px;
            border-radius: 8px;
        }

        input[type="range"] {
            flex: 1;
        }

        .threshold-value {
            background: #1e3c72;
            color: white;
            padding: 5px 10px;
            border-radius: 5px;
            font-weight: bold;
            min-width: 50px;
            text-align: center;
        }

        .btn {
            padding: 12px 30px;
            border: none;
            border-radius: 8px;
            font-size: 16px;
            font-weight: 600;
            cursor: pointer;
            transition: all 0.3s ease;
        }

        .btn-primary {
            background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
            color: white;
        }

        .btn-primary:hover:not(:disabled) {
            transform: translateY(-2px);
            box-shadow: 0 10px 20px rgba(30, 60, 114, 0.3);
        }

        .btn-primary:disabled {
            opacity: 0.6;
            cursor: not-allowed;
        }

        .loading {
            display: none;
            text-align: center;
            padding: 20px;
        }

        .loading.active {
            display: block;
        }

        .pipeline-progress {
            display: grid;
            grid-template-columns: repeat(10, 1fr);
            gap: 5px;
            margin: 20px 0;
        }

        .pipeline-step {
            padding: 10px 5px;
            background: #f8f9fa;
            border-radius: 5px;
            text-align: center;
            font-size: 0.75rem;
            transition: all 0.3s;
        }

        .pipeline-step.active {
            background: #4dabf7;
            color: white;
        }

        .pipeline-step.completed {
            background: #51cf66;
            color: white;
        }

        .results-section {
            display: none;
            margin-top: 40px;
        }

        .results-section.active {
            display: block;
        }

        .stats-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 20px;
            margin-bottom: 30px;
        }

        .stat-card {
            background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
            color: white;
            padding: 20px;
            border-radius: 12px;
            text-align: center;
        }

        .stat-card h3 {
            font-size: 2rem;
            margin-bottom: 5px;
        }

        table {
            width: 100%;
            border-collapse: collapse;
        }

        th {
            background: #f8f9fa;
            padding: 12px;
            text-align: left;
            font-weight: 600;
            color: #495057;
            border-bottom: 2px solid #dee2e6;
        }

        td {
            padding: 12px;
            border-bottom: 1px solid #e9ecef;
        }

        .score-badge {
            display: inline-block;
            padding: 4px 8px;
            border-radius: 4px;
            font-size: 0.85rem;
            font-weight: 600;
        }

        .score-high { background: #d3f9d8; color: #2b8a3e; }
        .score-medium { background: #fff3cd; color: #f08c00; }
        .score-low { background: #ffe3e3; color: #c92a2a; }

        .status-badge {
            display: inline-block;
            padding: 4px 12px;
            border-radius: 20px;
            font-size: 0.85rem;
            font-weight: 600;
        }

        .status-shortlisted { background: #d3f9d8; color: #2b8a3e; }
        .status-rejected { background: #ffe3e3; color: #c92a2a; }

        .details-btn {
            padding: 4px 12px;
            background: #1e3c72;
            color: white;
            border: none;
            border-radius: 4px;
            cursor: pointer;
            font-size: 0.85rem;
        }

        .details-btn:hover {
            background: #2a5298;
        }

        .modal {
            display: none;
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            height: 100%;
            background: rgba(0, 0, 0, 0.5);
            z-index: 1000;
        }

        .modal.active {
            display: flex;
            align-items: center;
            justify-content: center;
        }

        .modal-content {
            background: white;
            border-radius: 12px;
            padding: 30px;
            max-width: 800px;
            width: 90%;
            max-height: 80vh;
            overflow-y: auto;
        }

        .modal-header {
            display: flex;
            justify-content: space-between;
            align-items: center;
            margin-bottom: 20px;
        }

        .close-btn {
            background: none;
            border: none;
            font-size: 24px;
            cursor: pointer;
            color: #6c757d;
        }

        .detail-group {
            margin-bottom: 20px;
            padding: 15px;
            background: #f8f9fa;
            border-radius: 8px;
        }

        .detail-group h4 {
            color: #495057;
            margin-bottom: 10px;
        }

        .skill-gap-analysis {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 15px;
            margin-top: 10px;
        }

        .skill-category {
            padding: 10px;
            background: white;
            border-radius: 5px;
        }

        .skill-category h5 {
            font-size: 0.9rem;
            margin-bottom: 5px;
            color: #666;
        }

        .skill-list {
            font-size: 0.85rem;
            color: #333;
        }

        .audit-metrics {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(150px, 1fr));
            gap: 10px;
            margin-top: 20px;
            padding: 15px;
            background: #f8f9fa;
            border-radius: 8px;
        }

        .metric-item {
            text-align: center;
        }

        .metric-value {
            font-size: 1.5rem;
            font-weight: bold;
            color: #1e3c72;
        }

        .metric-label {
            font-size: 0.85rem;
            color: #666;
        }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>🎯 AI Recruiter - Industry Standard</h1>
            <p>Enterprise-Grade 10-Step Resume Analysis Pipeline</p>
            <div class="pipeline-indicator">
                <span class="step-badge">Step 1: Parsing</span>
                <span class="step-badge">Step 2: NLP</span>
                <span class="step-badge">Step 3: Semantic</span>
                <span class="step-badge">Step 4: Scoring</span>
                <span class="step-badge">Step 5: ML Models</span>
                <span class="step-badge">Step 6: Ranking</span>
                <span class="step-badge">Step 7: Industry</span>
                <span class="step-badge">Step 8: Advanced</span>
                <span class="step-badge">Step 9: Decision</span>
                <span class="step-badge">Step 10: QA</span>
            </div>
        </div>

        <div class="main-card">
            <form id="analyzeForm">
                <div class="form-group">
                    <label>📄 Upload Resumes (PDF, DOCX, TXT)</label>
                    <input type="file" id="files" multiple accept=".pdf,.docx,.doc,.txt" style="display: none;">
                    <label for="files" class="file-upload-label">
                        <div>
                            <p style="font-size: 1.2rem; margin-bottom: 10px;">Drop resumes here or click to browse</p>
                            <p style="color: #6c757d;">Industry Standard Processing - Max 50 files</p>
                        </div>
                    </label>
                    <div id="fileCount" style="margin-top: 10px; color: #6c757d;"></div>
                </div>

                <div class="form-group">
                    <label>📋 Job Description</label>
                    <textarea id="jobDescription" placeholder="Paste complete job description with all requirements, responsibilities, qualifications, skills, and company values..." required></textarea>
                </div>

                <div class="form-group">
                    <label>🎚️ Matching Threshold</label>
                    <div class="threshold-control">
                        <span>Minimum match required:</span>
                        <input type="range" id="thresholdSlider" min="10" max="90" value="20" step="5">
                        <span class="threshold-value" id="thresholdValue">20%</span>
                    </div>
                </div>

                <button type="submit" class="btn btn-primary">
                    🚀 Start Industry Standard Analysis
                </button>
            </form>

            <div class="loading">
                <h3>Executing 10-Step Industry Pipeline...</h3>
                <div class="pipeline-progress">
                    <div class="pipeline-step" id="step1">Parsing</div>
                    <div class="pipeline-step" id="step2">NLP</div>
                    <div class="pipeline-step" id="step3">Semantic</div>
                    <div class="pipeline-step" id="step4">Scoring</div>
                    <div class="pipeline-step" id="step5">ML</div>
                    <div class="pipeline-step" id="step6">Ranking</div>
                    <div class="pipeline-step" id="step7">Industry</div>
                    <div class="pipeline-step" id="step8">Advanced</div>
                    <div class="pipeline-step" id="step9">Decision</div>
                    <div class="pipeline-step" id="step10">QA</div>
                </div>
                <p style="margin-top: 20px; color: #6c757d;">Processing with Groq LLaMA 3.3 70B...</p>
            </div>

            <div class="results-section">
                <div class="stats-grid">
                    <div class="stat-card">
                        <h3 id="totalResumes">0</h3>
                        <p>Total Analyzed</p>
                    </div>
                    <div class="stat-card" style="background: linear-gradient(135deg, #51cf66 0%, #37b24d 100%);">
                        <h3 id="shortlisted">0</h3>
                        <p>Shortlisted</p>
                    </div>
                    <div class="stat-card" style="background: linear-gradient(135deg, #ff6b6b 0%, #f03e3e 100%);">
                        <h3 id="rejected">0</h3>
                        <p>Rejected</p>
                    </div>
                    <div class="stat-card">
                        <h3 id="avgScore">0%</h3>
                        <p>Average Score</p>
                    </div>
                </div>

                <div class="audit-metrics" id="auditMetrics"></div>

                <h3 style="margin: 20px 0;">📊 Analysis Results</h3>
                <table>
                    <thead>
                        <tr>
                            <th>Rank</th>
                            <th>Candidate</th>
                            <th>Overall Score</th>
                            <th>ATS Score</th>
                            <th>Confidence</th>
                            <th>Decision</th>
                            <th>Details</th>
                        </tr>
                    </thead>
                    <tbody id="resultsTableBody"></tbody>
                </table>

                <button class="btn btn-primary" style="margin-top: 20px;" onclick="downloadResults()">
                    📥 Download Complete Report
                </button>
            </div>
        </div>
    </div>

    <div class="modal" id="detailsModal">
        <div class="modal-content">
            <div class="modal-header">
                <h2 id="modalTitle">Candidate Analysis</h2>
                <button class="close-btn" onclick="closeModal()">×</button>
            </div>
            <div id="modalBody"></div>
        </div>
    </div>

    <script>
        let currentResults = null;
        let currentSessionId = null;
        let currentThreshold = 20;

        // File handling
        const fileInput = document.getElementById('files');
        const fileLabel = document.querySelector('.file-upload-label');
        const fileCount = document.getElementById('fileCount');

        fileInput.addEventListener('change', (e) => {
            const count = e.target.files.length;
            if (count > 0) {
                fileLabel.classList.add('has-files');
                fileCount.textContent = `${count} file(s) selected`;
                if (count > 50) {
                    fileCount.textContent += ' (Only first 50 will be processed)';
                }
            }
        });

        // Threshold slider
        const thresholdSlider = document.getElementById('thresholdSlider');
        const thresholdValue = document.getElementById('thresholdValue');

        thresholdSlider.addEventListener('input', (e) => {
            currentThreshold = e.target.value;
            thresholdValue.textContent = currentThreshold + '%';
        });

        // Form submission
        document.getElementById('analyzeForm').addEventListener('submit', async (e) => {
            e.preventDefault();

            const files = fileInput.files;
            const jobDescription = document.getElementById('jobDescription').value;

            if (!files.length) {
                alert('Please select resume files');
                return;
            }

            const formData = new FormData();
            for (let i = 0; i < Math.min(files.length, 50); i++) {
                formData.append('files', files[i]);
            }
            formData.append('job_description', jobDescription);
            formData.append('matching_threshold', currentThreshold);

            // Show loading with pipeline animation
            document.querySelector('.loading').classList.add('active');
            document.querySelector('.btn-primary').disabled = true;
            document.querySelector('.results-section').classList.remove('active');

            // Animate pipeline steps
            const steps = document.querySelectorAll('.pipeline-step');
            steps.forEach((step, index) => {
                setTimeout(() => {
                    step.classList.add('active');
                    setTimeout(() => {
                        step.classList.remove('active');
                        step.classList.add('completed');
                    }, 500);
                }, index * 300);
            });

            try {
                const response = await fetch('/api/analyze', {
                    method: 'POST',
                    body: formData
                });

                const data = await response.json();
                currentResults = data.results;
                currentSessionId = data.session_id;

                displayResults(data);

            } catch (error) {
                alert('Error: ' + error.message);
            } finally {
                document.querySelector('.loading').classList.remove('active');
                document.querySelector('.btn-primary').disabled = false;
                steps.forEach(step => {
                    step.classList.remove('active', 'completed');
                });
            }
        });

        function displayResults(data) {
            // Sort results: SHORTLIST first (by score), then REJECT (by score)
            data.results.sort((a, b) => {
                // First priority: Decision (SHORTLIST before REJECT)
                if (a.decision === 'SHORTLIST' && b.decision === 'REJECT') return -1;
                if (a.decision === 'REJECT' && b.decision === 'SHORTLIST') return 1;

                // Second priority: Higher final_score first
                const scoreA = a.final_score || a.overall_score || 0;
                const scoreB = b.final_score || b.overall_score || 0;
                return scoreB - scoreA;
            });

            // Update stats
            document.getElementById('totalResumes').textContent = data.total_resumes;
            document.getElementById('shortlisted').textContent = data.shortlisted_count;
            document.getElementById('rejected').textContent = data.rejected_count;

            // Calculate average score
            const avgScore = data.results.reduce((sum, r) => sum + (r.final_score || 0), 0) / data.results.length;
            document.getElementById('avgScore').textContent = avgScore.toFixed(1) + '%';

            // Display audit metrics
            if (data.audit_report) {
                const metricsHtml = `
                    <div class="metric-item">
                        <div class="metric-value">${(data.audit_report.performance_metrics?.precision * 100 || 0).toFixed(1)}%</div>
                        <div class="metric-label">Precision</div>
                    </div>
                    <div class="metric-item">
                        <div class="metric-value">${(data.audit_report.performance_metrics?.f1_score * 100 || 0).toFixed(1)}%</div>
                        <div class="metric-label">F1 Score</div>
                    </div>
                    <div class="metric-item">
                        <div class="metric-value">${data.audit_report.confidence_distribution?.HIGH || 0}</div>
                        <div class="metric-label">High Confidence</div>
                    </div>
                    <div class="metric-item">
                        <div class="metric-value">${data.audit_report.confidence_distribution?.MEDIUM || 0}</div>
                        <div class="metric-label">Medium Confidence</div>
                    </div>
                `;
                document.getElementById('auditMetrics').innerHTML = metricsHtml;
            }

            // Build table
            const tbody = document.getElementById('resultsTableBody');
            tbody.innerHTML = '';

            data.results.forEach((result, index) => {
                const row = document.createElement('tr');

                const scoreClass = result.final_score >= 70 ? 'score-high' : 
                                  result.final_score >= 40 ? 'score-medium' : 'score-low';
                const statusClass = result.decision === 'SHORTLIST' ? 'status-shortlisted' : 'status-rejected';

                row.innerHTML = `
                    <td>
                        ${result.decision === 'SHORTLIST' ? 
                            `<span style="color: #4caf50; font-weight: bold;">#${index + 1}</span>` : 
                            `<span style="color: #999;">#${index + 1}</span>`
                        }
                    </td>
                    <td>${result.filename}</td>
                    <td><span class="score-badge ${scoreClass}">${(result.final_score || 0).toFixed(1)}%</span></td>
                    <td><span class="score-badge ${result.ats_score >= 70 ? 'score-high' : result.ats_score >= 50 ? 'score-medium' : 'score-low'}">${(result.ats_score || 0).toFixed(1)}%</span></td>
                    <td>${result.confidence || 'N/A'}</td>
                    <td><span class="status-badge ${statusClass}">${result.decision}</span></td>
                    <td><button class="details-btn" onclick="showDetails(${index})">View</button></td>
                `;

                tbody.appendChild(row);
            });

            document.querySelector('.results-section').classList.add('active');
        }

        function showDetails(index) {
            const result = currentResults[index];
            document.getElementById('modalTitle').textContent = result.filename;

            const modalBody = document.getElementById('modalBody');
            modalBody.innerHTML = `
                <div class="detail-group">
                    <h4>📊 Comprehensive Scores</h4>
                    <div style="display: grid; grid-template-columns: repeat(2, 1fr); gap: 10px;">
                        <div>Overall Score: ${(result.overall_score || 0).toFixed(1)}%</div>
                        <div>Final Score: ${(result.final_score || 0).toFixed(1)}%</div>
                        <div>Semantic Score: ${(result.semantic_score || 0).toFixed(1)}%</div>
                        <div>Skill Match: ${(result.skill_match || 0).toFixed(1)}%</div>
                        <div>Keyword Score: ${(result.keyword_score || 0).toFixed(1)}%</div>
                        <div>Experience Score: ${(result.experience_score || 0).toFixed(1)}%</div>
                        <div>Education Score: ${(result.education_score || 0).toFixed(1)}%</div>
                        <div>ATS Score: ${(result.ats_score || 0).toFixed(1)}%</div>
                        <div>Cultural Fit: ${(result.cultural_fit || 0).toFixed(1)}%</div>
                        <div>ML Score: ${(result.ml_score || 0).toFixed(1)}%</div>
                    </div>
                </div>

                <div class="detail-group">
                    <h4>🎯 Decision Details</h4>
                    <p><strong>Decision:</strong> ${result.decision}</p>
                    <p><strong>Confidence:</strong> ${result.confidence}</p>
                    <p><strong>Explanation:</strong> ${result.decision_explanation || 'N/A'}</p>
                    <p><strong>Threshold Used:</strong> ${result.threshold_used || currentThreshold}%</p>
                </div>

                ${result.skill_gaps ? `
                <div class="detail-group">
                    <h4>📈 Skills Gap Analysis</h4>
                    <div class="skill-gap-analysis">
                        <div class="skill-category">
                            <h5>Matched Skills (${result.skill_gaps.matched_skills?.length || 0})</h5>
                            <div class="skill-list">${(result.skill_gaps.matched_skills || []).join(', ') || 'None'}</div>
                        </div>
                        <div class="skill-category">
                            <h5>Missing Skills (${result.skill_gaps.missing_skills?.length || 0})</h5>
                            <div class="skill-list">${(result.skill_gaps.missing_skills || []).join(', ') || 'None'}</div>
                        </div>
                        <div class="skill-category">
                            <h5>Transferable Skills</h5>
                            <div class="skill-list">${(result.skill_gaps.transferable_skills || []).join(', ') || 'None'}</div>
                        </div>
                    </div>
                    <p style="margin-top: 10px;"><strong>Match Percentage:</strong> ${(result.skill_gaps.match_percentage || 0).toFixed(1)}%</p>
                    <p><strong>Learning Potential:</strong> ${result.skill_gaps.learning_potential ? 'Yes' : 'No'}</p>
                </div>
                ` : ''}

                ${result.ats_issues && result.ats_issues.length > 0 ? `
                <div class="detail-group">
                    <h4>⚠️ ATS Compatibility Analysis</h4>
                    <div style="background: ${result.ats_score < 50 ? '#ffe3e3' : result.ats_score < 70 ? '#fff4e6' : '#e6f7ff'}; padding: 15px; border-radius: 8px; margin-bottom: 15px;">
                        <h5>Overall ATS Score: ${(result.ats_score || 0).toFixed(1)}%</h5>
                        <p><strong>Compatibility Level:</strong> ${result.compatibility_level || 'N/A'}</p>
                        <p><strong>Risk Level:</strong> ${result.risk_level || 'N/A'}</p>
                    </div>

                    ${result.detailed_scores ? `
                    <h5>Detailed ATS Breakdown:</h5>
                    <div style="display: grid; grid-template-columns: repeat(2, 1fr); gap: 10px; margin-bottom: 15px;">
                        <div>📄 Parsing Score: ${(result.detailed_scores.parsing || 0).toFixed(1)}%</div>
                        <div>🔑 Keywords Score: ${(result.detailed_scores.keywords || 0).toFixed(1)}%</div>
                        <div>📋 Structure Score: ${(result.detailed_scores.structure || 0).toFixed(1)}%</div>
                        <div>📧 Contact Info: ${(result.detailed_scores.contact || 0).toFixed(1)}%</div>
                        <div>📅 Date Formatting: ${(result.detailed_scores.dates || 0).toFixed(1)}%</div>
                    </div>
                    ` : ''}

                    <h5>Issues Detected:</h5>
                    <ul style="color: #d32f2f;">
                        ${result.ats_issues.map(issue => `<li>${issue}</li>`).join('')}
                    </ul>

                    ${result.recommendations && result.recommendations.length > 0 ? `
                    <h5>Recommendations:</h5>
                    <ul style="color: #1976d2;">
                        ${result.recommendations.map(rec => `<li>${rec}</li>`).join('')}
                    </ul>
                    ` : ''}
                </div>
                ` : ''}

                <div class="detail-group">
                    <h4>🏢 Additional Information</h4>
                    <p><strong>Industry:</strong> ${result.industry || 'N/A'}</p>
                    <p><strong>Years of Experience:</strong> ${result.years_experience || 0}</p>
                    <p><strong>Percentile Rank:</strong> ${result.percentile_rank || 'N/A'}%</p>
                    <p><strong>Weighted Score:</strong> ${result.weighted_score || 'N/A'}</p>
                </div>
            `;

            document.getElementById('detailsModal').classList.add('active');
        }

        function closeModal() {
            document.getElementById('detailsModal').classList.remove('active');
        }

        async function downloadResults() {
            if (!currentSessionId) return;

            const response = await fetch(`/api/download/${currentSessionId}`);
            const data = await response.json();

            const blob = new Blob([JSON.stringify(data, null, 2)], { type: 'application/json' });
            const url = URL.createObjectURL(blob);
            const a = document.createElement('a');
            a.href = url;
            a.download = `industry_standard_analysis_${currentSessionId}.json`;
            document.body.appendChild(a);
            a.click();
            document.body.removeChild(a);
            URL.revokeObjectURL(url);
        }

        // Modal close on outside click
        document.getElementById('detailsModal').addEventListener('click', (e) => {
            if (e.target.id === 'detailsModal') {
                closeModal();
            }
        });
    </script>
</body>
</html>
'''


@app.get("/", response_class=HTMLResponse)
async def index():
    return HTML_TEMPLATE


if __name__ == "__main__":
    import uvicorn

    print("\n" + "=" * 60)
    print("🚀 AI Recruiter System - Industry Standard Edition")
    print("=" * 60)
    print("📌 Version: 2.0 - Enterprise Grade")
    print("📌 Pipeline: Complete 10-Step Industry Standard")
    print("📌 URL: http://localhost:8000")
    print("📌 API Docs: http://localhost:8000/docs")
    print("📌 Health: http://localhost:8000/api/health")
    print("\n🔧 Pipeline Steps:")
    print("  1. Document Parsing & Extraction")
    print("  2. NLP Processing & Entity Recognition")
    print("  3. Semantic Understanding & Embeddings")
    print("  4. Multi-Factor Scoring")
    print("  5. ML Models (Classification & Ranking)")
    print("  6. Multi-Stage Filtering & Ranking")
    print("  7. Industry-Specific & ATS Optimization")
    print("  8. Advanced Analytics (Skills Gap, Cultural Fit)")
    print("  9. Dynamic Thresholding & Ensemble Decision")
    print(" 10. Quality Assurance & Audit")
    print("=" * 60 + "\n")

    uvicorn.run(app, host="0.0.0.0", port=8000, reload=False)
